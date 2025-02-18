
from flask import Blueprint, render_template, redirect, url_for, request, flash, send_file, g, session
from extensions import db
import os
from forms import (
    TeacherForm, ClassForm, GradeForm, AttendanceForm, StudentForm, AssignTeachersForm,
    AssessmentForm, AssessmentResultForm, SubjectForm, SchoolForm, UserForm, AssessmentTypeForm, AssignSubjectToClassForm, LoginForm
)
from models import (
    User, Student, Teacher, Class, Attendance, Assessment, AssessmentType,
    AssessmentSubjectScore, AssessmentResult, Subject, ClassSubject, TeacherSubject, School
)
from models import Grade, FeeComponent, StudentClassFeePayment, ClassFeeComponent
from flask_login import login_user, logout_user, login_required, current_user
from io import BytesIO
import pandas as pd
from docx import Document
from flask import send_from_directory
from sqlalchemy.sql import func
from werkzeug.utils import secure_filename
from flask import current_app
from config import Config
from sqlalchemy.orm import joinedload
from reportlab.lib.pagesizes import A4, letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate
from io import BytesIO
from flask import send_file, jsonify
import json
from collections import defaultdict
from sqlalchemy.orm import aliased

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from flask import g 
from utils import school_required, filter_by_school
import uuid


# Create a Blueprint instance
main = Blueprint('main', __name__)

    

@main.route('/static/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(current_app.config['UPLOAD_FOLDER'], filename)
    

@main.before_request
def load_school_details():
    if current_user.is_authenticated:
        # Check if school details exist in session
        if 'school_name' not in session or 'school_logo' not in session:
            school = School.query.get(current_user.school_id)
            if school:
                g.school_name = school.name
                g.school_logo = school.school_logo if school.school_logo else 'static/uploads/default-logo.png'
                session['school_name'] = g.school_name
                session['school_logo'] = g.school_logo  # Store in session
        else:
            g.school_name = session['school_name']
            g.school_logo = session['school_logo']
    else:
        g.school_name = "School Management Portal"
        g.school_logo = "uploads/default-logo.png"  


@main.route('/', methods=['GET', 'POST'])
def index():
    form = LoginForm()  
    if form.validate_on_submit():  
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user)

            # Store school details in session & `g`
            if user.school_id:
                school = School.query.get(user.school_id)
                if school:
                    g.school_name = school.name
                    g.school_logo = school.school_logo if school.school_logo else 'uploads/default-logo.png'
                    session['school_name'] = g.school_name
                    session['school_logo'] = g.school_logo  # Store in session

            # Redirect based on role
            next_page = request.args.get('next')  
            if next_page:
                return redirect(next_page)

            if user.role == 'admin':
                return redirect(url_for('main.dashboard'))
            elif user.role == 'teacher':
                return redirect(url_for('teacher.dashboard'))
            else:  
                return redirect(url_for('student.dashboard'))

        flash('Invalid username or password', 'danger')

    return render_template('index.html', form=form, is_authenticated=current_user.is_authenticated)




def generate_unique_code():
    return str(uuid.uuid4())

# When you create a new school:
# new_school = School(name="My School", registration_code=generate_unique_code())
# db.session.add(new_school)
# db.session.commit()

# @main.route('/register_school', methods=['GET', 'POST'])
# def register_school():
#     form = SchoolForm()
#     if form.validate_on_submit():
#         school_logo_path = None

#         if form.school_logo.data:
#             logo_file = form.school_logo.data
#             filename = secure_filename(logo_file.filename)
#             upload_folder = current_app.config['UPLOAD_FOLDER']  # Correct way to access config
#             filepath = os.path.join(upload_folder, filename)

#             try:
#                 logo_file.save(filepath)
#                 school_logo_path = f'static/images/school_logos/{filename}'  # Store relative path
#                 flash('Logo uploaded successfully!', 'success')
#             except Exception as e:
#                 flash(f'Error uploading logo: {e}', 'danger')
#                 print(f"Logo upload error: {e}")

#         school = School(
#             name=form.name.data,
#             address=form.address.data,
#             email=form.email.data,
#             phone_number=form.phone_number.data,
#             school_logo=school_logo_path,  # Store relative path
#             website=form.website.data
#         )

#         try:
#             db.session.add(school)
#             db.session.commit()
#             flash('School registered successfully!', 'success')
#             return redirect(url_for('main.register_user'))
#         except Exception as e:
#             db.session.rollback()
#             flash(f"Database error: {e}", "danger")
#             print(f"Database error: {e}")

#     return render_template('register_school.html', form=form)

@main.route('/register_school', methods=['GET', 'POST'])
def register_school():
    form = SchoolForm()
    if form.validate_on_submit():
        school_logo_path = None

        if form.school_logo.data:
            logo_file = form.school_logo.data
            filename = secure_filename(logo_file.filename)  # Sanitize filename
            upload_folder = current_app.config['UPLOAD_FOLDER']
            filepath = os.path.join(upload_folder, filename)

            # Check if the upload folder exists; create it if it doesn't
            os.makedirs(upload_folder, exist_ok=True)  # Important!

            try:
                logo_file.save(filepath)
                school_logo_path = f'static/images/school_logos/{filename}'
                flash('Logo uploaded successfully!', 'success')
            except Exception as e:
                flash(f'Error uploading logo: {e}', 'danger')
                print(f"Logo upload error: {e}")
                # Consider logging the full traceback for debugging:
                import traceback
                traceback.print_exc()  # Print the traceback to the console

        registration_code = str(uuid.uuid4())  # Generate a unique registration code
        school = School(
            name=form.name.data,
            address=form.address.data,
            email=form.email.data,
            phone_number=form.phone_number.data,
            school_logo=school_logo_path,
            website=form.website.data,
            registration_code=registration_code  # Add the registration code
        )

        try:
            db.session.add(school)
            db.session.commit()
            flash('School registered successfully!  Provide this registration link to the school: ' + 
                  url_for('main.register_user', school_code=registration_code, _external=True), 'success') # Show the link!

            # Redirect to a confirmation page or the school's dashboard
            return redirect(url_for('main.school_confirmation', school_id=school.id)) # Example redirect
        except Exception as e:
            db.session.rollback()
            flash(f"Database error: {e}", "danger")
            print(f"Database error: {e}")
            import traceback
            traceback.print_exc()


    return render_template('register_school.html', form=form)


@main.route('/school_confirmation/<int:school_id>')
def school_confirmation(school_id):
    school = School.query.get_or_404(school_id)
    return render_template('school_confirmation.html', school=school)


@main.route('/schools')
def schools():
    schools = School.query.all()
    return render_template('schools.html', schools=schools)

# Route to delete a school
@main.route('/delete/<int:id>')
def delete_school(id):
    school = School.query.get_or_404(id)
    db.session.delete(school)
    db.session.commit()
    return redirect(url_for('main.schools'))


@main.route('/edit/<int:id>', methods=['GET', 'POST'])
def edit_school(id):
    school = School.query.get_or_404(id)
    form = SchoolForm(obj=school)

    photo_filename = None

    if form.validate_on_submit():
        print("Form is valid")
        print("Form Data:", form.data)

        school.name = form.name.data
        school.address = form.address.data
        school.email = form.email.data
        school.phone_number = form.phone_number.data
        school.website = form.website.data
        school.school_logo = photo_filename
    
    

        # Handle photo upload
        if form.school_logo.data:
            try:
                photo_file = form.school_logo.data
                photo_filename = secure_filename(photo_file.filename)
                photo_path = os.path.join(Config.UPLOAD_FOLDER, photo_filename)
                os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)

                # Delete old logo if it exists (and if you're storing the full path)
                if school.school_logo:
                    old_photo_path = os.path.join(Config.UPLOAD_FOLDER, school.school_logo)
                    if os.path.exists(old_photo_path):
                        os.remove(old_photo_path)

                photo_file.save(photo_path)
                school.school_logo = photo_filename

                flash('Logo updated successfully!', 'success')

            except Exception as e:
                flash(f'Error updating logo: {e}', 'danger')
                print(f"Logo update error: {e}")

        try:
            db.session.commit()
            flash('School updated successfully!', 'success')
            return redirect(url_for('main.schools'))
        except Exception as e:
            db.session.rollback()
            flash(f"Database error: {e}", "danger")
            print(f"Database error: {e}")
            return render_template('edit_schools.html', school=school, form=form)

    else:
        print("Form is not valid. Errors:", form.errors)

    return render_template('edit_schools.html', school=school, form=form)

@main.route('/register_user', methods=['GET', 'POST'])
def register_user():
    school_code = request.args.get('school_code') # Get the school code from the URL

    if not school_code:
        flash('Invalid registration link.', 'danger')
        return redirect(url_for('main.index')) # Or wherever you want to redirect

    school = School.query.filter_by(registration_code=school_code).first() # Assuming you add 'registration_code' to your School model
    if not school:
        flash('Invalid school code.', 'danger')
        return redirect(url_for('main.index'))
    
    form = UserForm()
    form.school_id.choices = [(school.id, school.name)] # Limit choices to the specific school
    form.school_id.data = school.id # Pre-select the school (important!)
    form.school_id.render_kw = {'readonly': True} # Make the school field read-only in the form

    if form.validate_on_submit():
        if User.query.filter_by(username=form.username.data).first():
            flash('Username already exists!', 'danger')
            return redirect(url_for('main.register_user'))
        user = User(
            username=form.username.data,
            role=form.role.data,
            school_id=school.id # Use the school ID directly
        )
        user.set_password(form.password.data)
        try:
            db.session.add(user)
            db.session.commit()
            flash('User registered successfully!', 'success')
            return redirect(url_for('main.index'))
        except Exception as e:
            db.session.rollback()
            flash('An error occurred while registering the user.', 'danger')
    return render_template('register_user.html', form=form, school_name=school.name)



@main.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'success')
    return redirect(url_for('main.index'))

# Student Dashboard (Only accessible by Students)
@main.route('/student/dashboard')
@login_required
def student_dashboard():
    if current_user.role!= 'student':
        return redirect(url_for('main.index'))
    student = Student.query.get(current_user.id).last_name
    return render_template('student_dashboard.html', student=student)

# Teacher Dashboard (Only accessible by Teachers)
@main.route('/teacher/dashboard')
@login_required
def teacher_dashboard():
    if current_user.role != 'teacher':
        return redirect(url_for('main.index'))
    teacher_name = Teacher.query.get(current_user.id).name
    return render_template('teacher_dashboard.html', teacher_name=teacher_name)



# Dashboard (Only accessible by Admins)
# @main.route('/dashboard')
# # @school_required
# # @filter_by_school

# @login_required
# def dashboard():
#     if not current_user.role== 'admin':
#         return redirect(url_for('main.index'))
#     school_name = School.query.get(current_user.school_id).name
#     student_count = Student.query.count()
#     teacher_count = Teacher.query.count()
#     class_count = Class.query.count()
#     return render_template('dashboard.html', school_name=school_name, student_count=student_count,
#                            teacher_count=teacher_count, class_count=class_count)


@main.route('/dashboard')
@login_required
def dashboard():
    if current_user.role != 'admin':
        return redirect(url_for('main.index'))

    # Get the school linked to the logged-in user
    school = School.query.get(current_user.school_id)
    if not school:
        flash("No school linked to your account.", "danger")
        return redirect(url_for('main.index'))

    school_name = school.name

    # Filter counts by school_id
    student_count = Student.query.filter_by(school_id=current_user.school_id).count()
    teacher_count = Teacher.query.filter_by(school_id=current_user.school_id).count()
    class_count = Class.query.filter_by(school_id=current_user.school_id).count()

    return render_template('dashboard.html', 
                           school_name=school_name, 
                           student_count=student_count,
                           teacher_count=teacher_count, 
                           class_count=class_count)


# Report Generation Route
@main.route('/generate_report/<report_type>')
@login_required
def generate_report(report_type):

    # Fetching data for report
    data = {
        "Students": pd.read_sql(Student.query.statement, db.session.bind),
        "Teachers": pd.read_sql(Teacher.query.statement, db.session.bind),
        "Classes": pd.read_sql(Class.query.statement, db.session.bind),
        "Grades": pd.read_sql(Grade.query.statement, db.session.bind),
        "Attendance": pd.read_sql(Attendance.query.statement, db.session.bind),
    }
    
    if report_type == "excel":
        with BytesIO() as output:
            writer = pd.ExcelWriter(output, engine='xlsxwriter')
            for sheet, df in data.items():
                df.to_excel(writer, sheet_name=sheet, index=False)
            writer.save()
            output.seek(0)
            return send_file(output, attachment_filename="report.xlsx", as_attachment=True)
    
    elif report_type == "word":
        doc = Document()
        for section, df in data.items():
            doc.add_heading(section, level=1)
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                hdr_cells[i].text = column
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, attachment_filename="report.docx", as_attachment=True)


# Grades routes
@main.route('/grades')
def grades():
    grades = Grade.query.all()
    return render_template('grades.html', grades=grades)


@main.route('/attendance')
@login_required
def attendance():
    # Query attendance records with student names
    attendance_records = (
        db.session.query(
            Attendance.student_id,
            (Student.first_name + " " + Student.last_name).label("student_name"),
            Attendance.days_present,
            Attendance.total_days_opened
        )
        .join(Student, Attendance.student_id == Student.id)
        .all()
    )

    return render_template('attendance.html', attendance_records=attendance_records)


@main.route('/add_attendance', methods=['GET', 'POST'])
@login_required
def add_attendance():
    form = AttendanceForm()

    # Fetch students and their associated class names
    students = (
        db.session.query(
            Student.id,
            Student.first_name,
            Student.last_name,
            Class.class_name,
            Class.id.label('class_id')
        )
        .join(Class, Student.class_id == Class.id)
        .all()
    )

    # Set choices for the student_id field
    if students:
        form.student_id.choices = [(s.id, f"{s.first_name} {s.last_name}") for s in students]
    else:
        form.student_id.choices = [("", "No Students Available")]

    if form.validate_on_submit():
        student_id = form.student_id.data
        days_present = form.days_present.data
        total_days_opened = form.total_days_opened.data

        # Fetch the student's class automatically
        student = Student.query.get(student_id)
        if not student:
            flash("Invalid student selected.", "danger")
            return redirect(url_for('main.add_attendance'))

        class_id = student.class_id  # Auto-detect class

        # Check if an attendance record exists for this student
        attendance = Attendance.query.filter_by(student_id=student_id, class_id=class_id).first()

        if attendance:
            # Update attendance
            attendance.days_present = days_present
            attendance.total_days_opened = total_days_opened
        else:
            # Create a new attendance record
            attendance = Attendance(
                student_id=student_id,
                class_id=class_id,
                days_present=days_present,
                total_days_opened=total_days_opened
            )
            db.session.add(attendance)

        db.session.commit()
        flash('Attendance record updated successfully!', 'success')
        return redirect(url_for('main.attendance'))

    return render_template('add_attendance.html', form=form, students=students)


@main.route('/students', methods=['GET'])
@login_required
def students():
    """Display students only linked to the logged-in user's school."""
    search_query = request.args.get('search', '').strip().lower()

    if not current_user.school_id:
        flash("You are not linked to any school.", "warning")
        return redirect(url_for('main.dashboard'))

    query = Student.query.filter_by(school_id=current_user.school_id) 

    if search_query:
        query = query.join(Class).filter(
            (Student.first_name.ilike(f"%{search_query}%")) |
            (Student.last_name.ilike(f"%{search_query}%")) |
            (Class.class_name.ilike(f"%{search_query}%"))
        )

    students = query.all()

    return render_template('students.html', students=students)


# @main.route('/add_student', methods=['GET', 'POST'])
# def add_student():
#     form = StudentForm()

#     # Populate the choices dynamically for class_id
#     form.class_id.choices = [(cls.id, cls.class_name) for cls in Class.query.all()]
#     if not form.class_id.choices:
#         form.class_id.choices = [(-1, 'No classes available')]

#     # Populate the choices dynamically for school_id
#     form.school_id.choices = [(sch.id, sch.name) for sch in School.query.all()]
#     if not form.school_id.choices:
#         form.school_id.choices = [(-1, 'No schools available')]

#     if form.validate_on_submit():
#         # Handle the form submission
#         try:
#             if form.class_id.data == -1:
#                 flash("Please select a valid class.", "warning")
#                 return render_template('add_student.html', form=form)

#             if form.school_id.data == -1:
#                 flash("Please select a valid school.", "warning")
#                 return render_template('add_student.html', form=form)

#             student = Student(
#                 first_name=form.first_name.data,
#                 last_name=form.last_name.data,
#                 date_of_birth=form.date_of_birth.data,
#                 enrollment_date=form.enrollment_date.data,
#                 gender=form.gender.data,
#                 grade_level=form.grade_level.data,
#                 contact_email=form.contact_email.data,
#                 school_id=current_user.school_id,
#                 class_id=form.class_id.data
#             )
#             db.session.add(student)
#             db.session.commit()
#             flash('Student added successfully!', 'success')
#             return redirect(url_for('main.students'))
#         except Exception as e:
#             db.session.rollback()
#             flash(f"An error occurred while adding the student: {str(e)}", "danger")

#     return render_template('add_student.html', form=form)

@main.route('/add_student', methods=['GET', 'POST'])
@login_required  # Ensure the user is logged in
def add_student():
    form = StudentForm()

    # Get the current user's school
    current_school = current_user.school

    if not current_school: # Handle the case where the user has no school
        flash("You are not associated with any school. Please contact an administrator.", "warning")
        return redirect(url_for('main.index')) # Or another appropriate route

    # Populate class choices for the current user's school ONLY
    form.class_id.choices = [(cls.id, cls.class_name) for cls in Class.query.filter_by(school_id=current_school.id).all()]
    if not form.class_id.choices:
        form.class_id.choices = [(-1, 'No classes available for your school')]  # More specific message

    # School should be pre-filled and disabled
    form.school_id.choices = [(current_school.id, current_school.name)]
    form.school_id.data = current_school.id  # Set the current school ID
    form.school_id.render_kw = {'readonly': True}  # Make the field read-only


    if form.validate_on_submit():
        if form.class_id.data == -1:
            flash("Please select a valid class.", "warning")
            return render_template('add_student.html', form=form)

        try:
            student = Student(
                first_name=form.first_name.data,
                last_name=form.last_name.data,
                date_of_birth=form.date_of_birth.data,
                enrollment_date=form.enrollment_date.data,
                gender=form.gender.data,
                grade_level=form.grade_level.data,
                contact_email=form.contact_email.data,
                school_id=current_school.id,  # Use the current user's school ID
                class_id=form.class_id.data
            )
            db.session.add(student)
            db.session.commit()
            flash('Student added successfully!', 'success')
            return redirect(url_for('main.students'))
        except Exception as e:
            db.session.rollback()
            flash(f"An error occurred while adding the student: {str(e)}", "danger")

    return render_template('add_student.html', form=form)



@main.route('/edit_student/<int:student_id>', methods=['GET', 'POST'])
@login_required
def edit_student(student_id):
    student = Student.query.get_or_404(student_id)

    if student.school_id != current_user.school_id:  # Check school association
        flash("You do not have permission to edit this student.", "danger")
        return redirect(url_for('main.students'))

    form = StudentForm(obj=student)  # Pre-populate form with student data

    form.class_id.choices = [(cls.id, cls.class_name) for cls in Class.query.filter_by(school_id=student.school_id).all()]
    if not form.class_id.choices:
        form.class_id.choices = [(-1, 'No classes available for your school')]

    form.school_id.choices = [(student.school_id, student.school.name)] # pre-select school and disable it
    form.school_id.data = student.school_id
    form.school_id.render_kw = {'readonly': True}


    if form.validate_on_submit():
        if form.class_id.data == -1:
            flash("Please select a valid class.", "warning")
            return render_template('edit_student.html', form=form)

        try:
            student.first_name = form.first_name.data
            student.last_name = form.last_name.data
            student.date_of_birth = form.date_of_birth.data
            student.enrollment_date = form.enrollment_date.data
            student.gender = form.gender.data
            student.grade_level = form.grade_level.data
            student.contact_email = form.contact_email.data
            student.class_id = form.class_id.data
            # No need to update school_id, as it's disabled
            db.session.commit()
            flash('Student updated successfully!', 'success')
            return redirect(url_for('main.students'))
        except Exception as e:
            db.session.rollback()
            flash(f"An error occurred while updating the student: {str(e)}", "danger")

    return render_template('edit_student.html', form=form, student=student)


@main.route('/delete_student/<int:student_id>', methods=['POST'])  # POST for security
@login_required
def delete_student(student_id):
    student = Student.query.get_or_404(student_id)

    if student.school_id != current_user.school_id:  # Check school association
        flash("You do not have permission to delete this student.", "danger")
        return redirect(url_for('main.students'))

    try:
        db.session.delete(student)
        db.session.commit()
        flash('Student deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f"An error occurred while deleting the student: {str(e)}", "danger")

    return redirect(url_for('main.students'))

# @main.route('/teachers', methods=['GET'])
# def teachers():
#     search_query = request.args.get('search', '').strip().lower()
    
#     if search_query:
#         teachers = Teacher.query.filter(
#             (Teacher.first_name.ilike(f"%{search_query}%")) |
#             (Teacher.last_name.ilike(f"%{search_query}%")) |
#             (Teacher.qualification.ilike(f"%{search_query}%")) |
#             (Teacher.teacher_subjects.any(Subject.name.ilike(f"%{search_query}%")))
#         ).all()
#     else:
#         teachers = Teacher.query.all()
#     return render_template('teachers.html', teachers=teachers)


@main.route('/teachers', methods=['GET'])
@login_required  # Ensure the user is logged in
def teachers():
    search_query = request.args.get('search', '').strip().lower()

    current_user_school = current_user.school

    if not current_user_school:
      return redirect(url_for('main.dashboard'))

    if search_query:
        teachers = Teacher.query.filter(
            (Teacher.school == current_user_school) &
            (
                (Teacher.first_name.ilike(f"%{search_query}%")) |
                (Teacher.last_name.ilike(f"%{search_query}%")) |
                (Teacher.qualification.ilike(f"%{search_query}%")) |
                (Teacher.teacher_subjects.any(Subject.name.ilike(f"%{search_query}%")))
            )
        ).all()
    else:
        teachers = Teacher.query.filter(Teacher.school == current_user_school).all()

    return render_template('teachers.html', teachers=teachers)

# @main.route('/add_teacher', methods=['GET', 'POST'])
# @login_required
# def add_teacher():
#     form = TeacherForm()

#     # Populate choices
#     form.school_id.choices = [(school.id, school.name) for school in School.query.all()]
#     form.subject.choices = [(subject.id, subject.name) for subject in Subject.query.all()]

#     if form.validate_on_submit():
#         # Handle photo upload
#         photo_filename = None
#         if form.photo.data:
#             photo_file = form.photo.data
#             photo_filename = secure_filename(photo_file.filename)
#             photo_path = os.path.join(Config.UPLOAD_FOLDER, photo_filename)
#             os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
#             photo_file.save(photo_path)

#         # Save the teacher's data
#         teacher = Teacher(
#             first_name=form.first_name.data,
#             last_name=form.last_name.data,
#             email=form.email.data,
#             hire_date=form.hire_date.data,
#             phone_number=form.phone_number.data,
#             qualification=form.qualification.data,
#             address=form.address.data,
#             date_of_birth=form.date_of_birth.data,
#             gender=form.gender.data,
#             photo=photo_filename,
#             school_id=form.school_id.data
#         )
#         db.session.add(teacher)
#         db.session.commit()

#         # Debugging: Check if the teacher is saved
#         print(f"Teacher added: {teacher}")

#         # Associate teacher with subjects
#         try:
#             for subject_id in form.subject.data:
#                 teacher_subject = TeacherSubject(teacher_id=teacher.id, subject_id=subject_id)
#                 db.session.add(teacher_subject)
#                 print(f"Added TeacherSubject: {teacher_subject}")
#             db.session.commit()
#         except Exception as e:
#             db.session.rollback()
#             print(f"Error adding TeacherSubject: {e}")
#             flash("Failed to add subjects for the teacher.")
#             return redirect(url_for('main.teachers'))

#         flash('Teacher added successfully with photo and details!')
#         return redirect(url_for('main.teachers'))

#     return render_template('add_teacher.html', form=form)


@main.route('/add_teacher', methods=['GET', 'POST'])
@login_required
def add_teacher():
    form = TeacherForm()

    current_school = current_user.school
    if not current_school:
        flash("You are not associated with any school.", "warning")
        return redirect(url_for('main.index'))  # Or appropriate route

    # Filter schools and subjects by the logged-in user's school
    form.school_id.choices = [(current_school.id, current_school.name)]
    form.school_id.data = current_school.id  # Set the current school ID
    form.school_id.render_kw = {'readonly': True}  # Make the field read-only

    form.subject.choices = [(subject.id, subject.name) for subject in Subject.query.filter_by(school_id=current_school.id).all()]
    if not form.subject.choices:
        form.subject.choices = [(-1, 'No subjects available for your school')]

    if form.validate_on_submit():
        if form.subject.data and form.subject.data[0] == -1:
            flash("Please select at least one valid subject.", "warning")
            return render_template('add_teacher.html', form=form)

        photo_filename = None
        if form.photo.data:
            photo_file = form.photo.data
            photo_filename = secure_filename(photo_file.filename)
            photo_path = os.path.join(Config.UPLOAD_FOLDER, photo_filename)
            os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)  # Create directory if it doesn't exist
            photo_file.save(photo_path)

        teacher = Teacher(
            first_name=form.first_name.data,
            last_name=form.last_name.data,
            email=form.email.data,
            hire_date=form.hire_date.data,
            phone_number=form.phone_number.data,
            qualification=form.qualification.data,
            address=form.address.data,
            date_of_birth=form.date_of_birth.data,
            gender=form.gender.data,
            photo=photo_filename,
            school_id=current_school.id  # Use current_user.school_id
        )

        try:
            db.session.add(teacher)
            db.session.commit()  # Commit *before* adding subjects to get teacher.id

            if form.subject.data and form.subject.data[0] != -1: # Check if the user selected a subject
                for subject_id in form.subject.data:
                    teacher_subject = TeacherSubject(teacher_id=teacher.id, subject_id=subject_id)
                    db.session.add(teacher_subject)
                db.session.commit()
        except Exception as e:
            db.session.rollback()
            flash(f"An error occurred: {e}", "danger")  # More informative error message
            return render_template('add_teacher.html', form=form) # Redirect to the same page to show the error
        
        flash('Teacher added successfully!')
        return redirect(url_for('main.teachers'))

    return render_template('add_teacher.html', form=form)


@main.route('/delete_teacher/<int:teacher_id>', methods=['POST'])
@login_required
def delete_teacher(teacher_id):
    teacher = Teacher.query.get_or_404(teacher_id)

    # Delete associated subjects
    TeacherSubject.query.filter_by(teacher_id=teacher.id).delete()

    # Delete the teacher record
    db.session.delete(teacher)
    db.session.commit()
    flash('Teacher deleted successfully!')
    return redirect(url_for('main.teachers'))


@main.route('/edit_teacher/<int:teacher_id>', methods=['GET', 'POST'])
@login_required
def edit_teacher(teacher_id):
    teacher = Teacher.query.get_or_404(teacher_id)
    form = TeacherForm(obj=teacher)

    # Populate choices
    form.school_id.choices = [(school.id, school.name) for school in School.query.all()]
    form.subject.choices = [(subject.id, subject.name) for subject in Subject.query.all()]

    if form.validate_on_submit():
        teacher.first_name = form.first_name.data
        teacher.last_name = form.last_name.data
        teacher.email = form.email.data
        teacher.hire_date = form.hire_date.data
        teacher.phone_number = form.phone_number.data
        teacher.qualification = form.qualification.data
        teacher.address = form.address.data
        teacher.date_of_birth = form.date_of_birth.data
        teacher.gender = form.gender.data
        teacher.school_id = form.school_id.data

        # Update subjects
        TeacherSubject.query.filter_by(teacher_id=teacher.id).delete()
        for subject_id in form.subject.data:
            teacher_subject = TeacherSubject(teacher_id=teacher.id, subject_id=subject_id)
            db.session.add(teacher_subject)

        db.session.commit()
        flash('Teacher updated successfully!')
        return redirect(url_for('main.teachers'))

    return render_template('edit_teacher.html', form=form, teacher=teacher)


@main.route('/add_class', methods=['GET', 'POST'])
@login_required
def add_class():
    form = ClassForm()
    teachers = Teacher.query.filter_by(school_id=current_user.school_id).all()

    # Populate form teacher_ids choices dynamically
    form.teacher_ids.choices = [(teacher.id, f"{teacher.first_name} {teacher.last_name}") for teacher in teachers]

    if form.validate_on_submit():
        # Create a new class instance
        new_class = Class(
            class_name=form.class_name.data,
            class_level=form.class_level.data,
            class_category=form.class_category.data,
            school_id=current_user.school_id,
        )

        # Add selected teachers to the class
        selected_teacher_ids = form.teacher_ids.data
        selected_teachers = Teacher.query.filter(Teacher.id.in_(selected_teacher_ids)).all()
        new_class.teachers.extend(selected_teachers)

        try:
            db.session.add(new_class)
            db.session.commit()
            flash('Class added successfully!', 'success')
            return redirect(url_for('main.classes'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error adding class: {str(e)}', 'danger')

    return render_template('add_class.html', form=form)


# @main.route('/classes')
# @login_required
# def classes():
#     # Query classes, join teachers, and calculate student counts
#     classes = db.session.query(
#         Class,
#         func.count(Student.id).label('student_count')
#     ).outerjoin(Student, Student.class_id == Class.id) \
#      .filter(Class.school_id == current_user.school_id) \
#      .group_by(Class.id).options(joinedload('teachers')).all()

#     return render_template('classes.html', classes=classes)


# @main.route('/edit_class/<int:class_id>', methods=['GET', 'POST'])
# def edit_class(class_id):
#     # Fetch the class by ID
#     class_to_edit = Class.query.get_or_404(class_id)

#     if request.method == 'POST':
#         # Update the class details
#         class_to_edit.class_name = request.form['class_name']
#         class_to_edit.class_category = request.form['class_category']
        
#         try:
#             db.session.commit()
#             flash('Class updated successfully!', 'success')
#             return redirect(url_for('main.classes'))
#         except Exception as e:
#             db.session.rollback()
#             flash(f'Error updating class: {str(e)}', 'danger')

#     return render_template('edit_class.html', class_to_edit=class_to_edit)

# @main.route('/delete_class/<int:class_id>', methods=['GET', 'POST'])
# def delete_class(class_id):
#     # Fetch the class by ID
#     class_to_delete = Class.query.get_or_404(class_id)

#     try:
#         db.session.delete(class_to_delete)
#         db.session.commit()
#         flash('Class deleted successfully!', 'success')
#     except Exception as e:
#         db.session.rollback()
#         flash(f'Error deleting class: {str(e)}', 'danger')

#     return redirect(url_for('main.classes'))


@main.route('/classes')
@login_required
def classes():
    classes_data = db.session.query(
        Class,
        func.count(Student.id).label('student_count')
    ).outerjoin(Student, Student.class_id == Class.id) \
     .filter(Class.school_id == current_user.school_id) \
     .group_by(Class.id).options(joinedload('teachers')).all()

    # Separate Class objects and student counts
    classes = []
    for class_obj, student_count in classes_data:
        classes.append({
            'class': class_obj,
            'student_count': student_count
        })

    return render_template('classes.html', classes=classes)

@main.route('/edit_class/<int:class_id>', methods=['GET', 'POST'])
@login_required  # Add login_required
def edit_class(class_id):
    class_to_edit = Class.query.get_or_404(class_id)

    if class_to_edit.school_id != current_user.school_id:  # Check school association
        flash("You do not have permission to edit this class.", "danger")
        return redirect(url_for('main.classes'))

    if request.method == 'POST':
        class_to_edit.class_name = request.form['class_name']
        class_to_edit.class_category = request.form['class_category']
        class_to_edit.class_level = request.form['class_level'] # Add this line
        try:
            db.session.commit()
            flash('Class updated successfully!', 'success')
            return redirect(url_for('main.classes'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating class: {str(e)}', 'danger')

    return render_template('edit_class.html', class_to_edit=class_to_edit)

@main.route('/delete_class/<int:class_id>', methods=['POST'])
@login_required
def delete_class(class_id):
    class_to_delete = Class.query.get_or_404(class_id)

    if class_to_delete.school_id != current_user.school_id:
        flash("You do not have permission to delete this class.", "danger")
        return redirect(url_for('main.classes'))

    try:
        # SQLAlchemy should handle the association table automatically
        db.session.delete(class_to_delete)  # Delete the class
        db.session.commit()
        flash('Class deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting class: {str(e)}', 'danger')

    return redirect(url_for('main.classes'))

@main.route('/add_grade', methods=['GET', 'POST'])
@login_required
def add_grade():
    form = GradeForm()
    if form.validate_on_submit():
        grade = Grade(
            student_id=form.student_id.data,
            class_id=form.class_id.data,
            score=form.score.data
        )
        db.session.add(grade)
        db.session.commit()
        flash('Grade added successfully!')
        return redirect(url_for('main.grades'))
    return render_template('add_grade.html', form=form)


# Route to display all subjects
@main.route('/subjects')
@login_required
def subjects():
    subjects = Subject.query.filter_by(school_id=current_user.school_id).all()
    return render_template('subjects.html', subjects=subjects)

# Route to add a subject
# @main.route('/add_subject', methods=['GET', 'POST'])
# @login_required
# def add_subject():
#     form = SubjectForm()
#     if form.validate_on_submit():
#         subject = Subject(
#             name=form.name.data,
#             school_id=current_user.school_id
#         )
#         db.session.add(subject)
#         db.session.commit()
#         flash('Subject added successfully!')
#         return redirect(url_for('main.subjects'))
#     return render_template('add_subject.html', form=form)

@main.route('/add_subject', methods=['GET', 'POST'])
@login_required
def add_subject():
    # Ensure the user is linked to a school
    if not current_user.school_id:
        flash("You are not linked to any school.", "danger")
        return redirect(url_for('main.index'))

    # Ensure only admins can add subjects
    if current_user.role != 'admin':
        flash("You are not authorized to add subjects.", "danger")
        return redirect(url_for('main.index'))

    form = SubjectForm()
    
    if form.validate_on_submit():
        # Ensure the subject does not already exist in the school
        existing_subject = Subject.query.filter_by(name=form.name.data, school_id=current_user.school_id).first()
        if existing_subject:
            flash("This subject already exists in your school.", "warning")
            return redirect(url_for('main.add_subject'))

        # Create and save the new subject
        subject = Subject(
            name=form.name.data.strip(),  # Remove accidental spaces
            school_id=current_user.school_id
        )
        db.session.add(subject)
        db.session.commit()
        flash('Subject added successfully!', 'success')
        return redirect(url_for('main.subjects'))

    return render_template('add_subject.html', form=form)

@main.route('/assign_subject_to_class', methods=['GET', 'POST'])
@login_required
def assign_subject_to_class():
    form = AssignSubjectToClassForm()

    # Populate the class dropdown
    form.class_id.choices = [
        (c.id, c.class_name)  # (value, label)
        for c in Class.query.filter_by(school_id=current_user.school_id).all()
    ]

    # Fetch all subjects for the current school
    subjects = Subject.query.filter_by(school_id=current_user.school_id).all()

    if form.validate_on_submit():
        class_id = form.class_id.data
        selected_subjects = request.form.getlist('subjects')

        # Assign each selected subject to the class
        for subject_id in selected_subjects:
            class_subject = ClassSubject(
                class_id=class_id,
                subject_id=int(subject_id)
            )
            db.session.add(class_subject)
        
        db.session.commit()
        flash('Subjects assigned to class successfully!', 'success')
        return redirect(url_for('main.subjects'))

    return render_template('assign_subject_to_class.html', form=form, subjects=subjects)


@main.route('/class_subjects')
def view_class_subjects():
    # Define table aliases to prevent ambiguity
    subject_alias = aliased(Subject)

    # Query data and join with classes & subjects
    class_subjects = db.session.query(
        Class.class_name.label("class_name"),
        subject_alias.name.label("subject_name")
    ).join(ClassSubject, Class.id == ClassSubject.class_id) \
     .join(subject_alias, subject_alias.id == ClassSubject.subject_id) \
     .all()

    # Group subjects by class name
    grouped_class_subjects = {}
    for cs in class_subjects:
        grouped_class_subjects.setdefault(cs.class_name, []).append(cs.subject_name)

    return render_template("view_class_subjects.html", grouped_class_subjects=grouped_class_subjects)

@main.route('/edit_class_subject/<class_name>', methods=['GET', 'POST'])
def edit_class_subject(class_name):
    # Fetch subjects linked to the class
    class_obj = Class.query.filter_by(class_name=class_name).first()
    if not class_obj:
        flash("Class not found", "danger")
        return redirect(url_for('main.view_class_subjects'))

    subjects = Subject.query.all()
    class_subjects = ClassSubject.query.filter_by(class_id=class_obj.id).all()
    selected_subjects = [cs.subject_id for cs in class_subjects]

    if request.method == 'POST':
        new_subjects = request.form.getlist('subjects')
        
        # Remove existing subjects not in new selection
        for cs in class_subjects:
            if str(cs.subject_id) not in new_subjects:
                db.session.delete(cs)
        
        # Add new selections
        for subject_id in new_subjects:
            if int(subject_id) not in selected_subjects:
                new_entry = ClassSubject(class_id=class_obj.id, subject_id=int(subject_id))
                db.session.add(new_entry)

        db.session.commit()
        flash("Subjects updated successfully!", "success")
        return redirect(url_for('main.view_class_subjects'))

    return render_template("edit_class_subject.html", class_obj=class_obj, subjects=subjects, selected_subjects=selected_subjects)

@main.route('/delete_class_subject/<class_name>', methods=['GET', 'POST'])
def delete_class_subject(class_name):
    class_obj = Class.query.filter_by(class_name=class_name).first()
    if not class_obj:
        flash("Class not found", "danger")
        return redirect(url_for('main.view_class_subjects'))

    ClassSubject.query.filter_by(class_id=class_obj.id).delete()
    db.session.commit()
    
    flash("Class subjects deleted successfully!", "success")
    return redirect(url_for('main.view_class_subjects'))


# Route to generate a report of all assessment scores
@main.route('/generate_assessment_report/<int:assessment_id>')
@login_required
def generate_assessment_report(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    scores = AssessmentResult.query.filter_by(assessment_id=assessment_id).all()
    data = pd.DataFrame([{
        "Student": f"{score.student.first_name} {score.student.last_name}",
        "Subject": score.subject.name,
        "Marks Obtained": score.marks_obtained
    } for score in scores])
    
    with BytesIO() as output:
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        data.to_excel(writer, index=False, sheet_name='Assessment Scores')
        writer.save()
        output.seek(0)
        return send_file(output, attachment_filename=f"Assessment_{assessment.name}_Report.xlsx", as_attachment=True)


# Route to list all assessment types
# @main.route('/assessment_types')
# @login_required
# def assessment_types():
#     if not current_user.role == 'admin':
#         flash('Access denied!', 'danger')
#         return redirect(url_for('main.index'))
#     assessment_types = AssessmentType.query.all()
#     return render_template('assessment_types.html', assessment_types=assessment_types)


@main.route('/assessment_types')
@login_required
def assessment_types():
    if not current_user.role == 'admin'and current_user.school_id is None:
        flash('Access denied!', 'danger')
        return redirect(url_for('main.index'))

    # Option 1: Direct Relationship (Preferred if exists)
    # Assumes you have a direct relationship like School.assessment_types
    # and AssessmentType has a school_id or school relationship.
    # if hasattr(current_user.school, 'assessment_types'): #Check if the relationship exists
    #     assessment_types = current_user.school.assessment_types
    # else:
    #     flash('No Assessment Types found for this school', 'info')
    #     assessment_types = []


    # Option 2: Join and Filter (If no direct relationship)
    # Use this if there's no direct School.assessment_types relationship
    # but AssessmentType has a school_id.
    assessment_types = AssessmentType.query.filter(AssessmentType.school_id == current_user.school_id).all()

    return render_template('assessment_types.html', assessment_types=assessment_types)

# Super Admin Use only
@main.route('/assessment_type')
@login_required
def assessment_type():
    if not current_user.role == 'admin':
        flash('Access denied!', 'danger')
        return redirect(url_for('main.index'))
    
    # Filter assessment types by the logged-in user's school
    assessment_types = AssessmentType.query.all()
    
    return render_template('assessment_types.html', assessment_types=assessment_types)

# Route to add an assessment type
# @main.route('/add_assessment_type', methods=['GET', 'POST'])
# @login_required
# def add_assessment_type():
#     form = AssessmentTypeForm()
#     if not current_user.role == 'admin':
#         flash('Access denied!', 'danger')
#         return redirect(url_for('main.index'))
#     if request.method == 'POST':
#         name = request.form.get('name')
#         if not name:
#             flash('Assessment type name is required.', 'danger')
#         else:
#             new_type = AssessmentType(name=name)
#             db.session.add(new_type)
#             db.session.commit()
#             flash('Assessment type added successfully!', 'success')
#             return redirect(url_for('main.assessment_types'))
#     return render_template('add_assessment_type.html', form=form)


# @main.route('/add_assessment_type', methods=['GET', 'POST'])
# @login_required
# def add_assessment_type():
#     form = AssessmentTypeForm()
#     if not current_user.role == 'admin':
#         flash('Access denied!', 'danger')
#         return redirect(url_for('main.index'))

#     if request.method == 'POST':
#         name = request.form.get('name')
#         school_id = request.form.get('school')

#         if not name:
#             flash('Assessment type name is required.', 'danger')
#         elif not school_id:
#             flash('School is required.', 'danger')
#         else:
#             try:
#                 # Convert school_id to integer (important!)
#                 school_id = int(school_id)
#                 new_type = AssessmentType(name=name, school_id=school_id)
#                 db.session.add(new_type)
#                 db.session.commit()
#                 flash('Assessment type added successfully!', 'success')
#                 return redirect(url_for('main.assessment_types'))
#             except ValueError:
#                 flash('Invalid school ID.', 'danger')
#             except Exception as e:
#                 db.session.rollback()
#                 flash(f'An error occurred: {str(e)}', 'danger')
                
#     return render_template('add_assessment_type.html', form=form)

# @main.route('/add_assessment_type', methods=['GET', 'POST'])
# @login_required
# def add_assessment_type():
#     # ... (rest of your code)
#     form = AssessmentTypeForm()
#     if not current_user.role == 'admin':
#         flash('Access denied!', 'danger')
#         return redirect(url_for('main.index'))

#     if request.method == 'POST':
#         name = request.form.get('name')
#         school_id = request.form.get('school')

#         # ... (validation for name and school_id)
#         if not name:
#             flash('Assessment type name is required.', 'danger')
#         elif not school_id:
#             flash('School is required.', 'danger')
#         else:

#             try:
#                 school_id = int(school_id)

#                 # Check if name already exists for this school
#                 existing_type = AssessmentType.query.filter_by(name=name, school_id=school_id).first()

#                 if existing_type:
#                     flash('An assessment type with that name already exists for this school.', 'danger')  # Specific error message
#                 else:
#                     new_type = AssessmentType(name=name, school_id=school_id)
#                     db.session.add(new_type)
#                     db.session.commit()
#                     flash('Assessment type added successfully!', 'success')
#                     return redirect(url_for('main.assessment_types'))

#             except ValueError:
#                 flash('Invalid school ID.', 'danger')
#             except Exception as e:
#                 db.session.rollback()
#                 flash(f'An error occurred: {str(e)}', 'danger')

#     return render_template('add_assessment_type.html', form=form)

# @main.route('/add_assessment_type', methods=['GET', 'POST'])
# @login_required
# def add_assessment_type():
#     form = AssessmentTypeForm()

#     # Restrict access to admins
#     if current_user.role != 'admin':
#         flash('Access denied!', 'danger')
#         return redirect(url_for('main.index'))

#     if request.method == 'POST':
#         name = request.form.get('name')

#         # Ensure name is provided
#         if not name:
#             flash('Assessment type name is required.', 'danger')
#         else:
#             try:
#                 # Automatically link to the logged-in user's school
#                 school_id = current_user.school_id  

#                 # Check if the assessment type already exists for this specific school
#                 existing_type = AssessmentType.query.filter_by(name=name, school_id=school_id).first()
#                 print(existing_type)

#                 if existing_type:
#                     flash('An assessment type with this name already exists for your school.', 'warning')
#                 else:
#                     new_type = AssessmentType(name=name.strip(), school_id=school_id)  # Trim spaces
#                     db.session.add(new_type)
#                     db.session.commit()
#                     flash('Assessment type added successfully!', 'success')
#                     return redirect(url_for('main.assessment_types'))

#             except Exception as e:
#                 db.session.rollback()
#                 flash(f'An error occurred: {str(e)}', 'danger')

#     return render_template('add_assessment_type.html', form=form)

@main.route('/add_assessment_type', methods=['GET', 'POST'])
@login_required
def add_assessment_type():
    form = AssessmentTypeForm()

    # Restrict access to admins
    if current_user.role != 'admin':
        flash('Access denied!', 'danger')
        return redirect(url_for('main.index'))

    if request.method == 'POST':
        name = request.form.get('name')

        if not name:
            flash('Assessment type name is required.', 'danger')
        else:
            try:
                school_id = current_user.school_id  # Automatically assign school

                # Check if the assessment type exists for this school
                existing_type = AssessmentType.query.filter_by(name=name, school_id=school_id).first()
                if existing_type:
                    flash('An assessment type with this name already exists for your school.', 'warning')
                else:
                    new_type = AssessmentType(name=name.strip(), school_id=school_id)
                    db.session.add(new_type)
                    db.session.commit()
                    flash('Assessment type added successfully!', 'success')
                    return redirect(url_for('main.assessment_types'))

            except Exception as e:
                db.session.rollback()
                flash(f'An error occurred: {str(e)}', 'danger')

    return render_template('add_assessment_type.html', form=form)


@main.route('/edit_assessment_type/<int:assessment_type_id>', methods=['GET', 'POST'])  # Add GET method
@login_required
def edit_assessment_type(assessment_type_id):
    if not current_user.role == 'admin':
        flash('Access denied!', 'danger')
        return redirect(url_for('main.index'))

    assessment_type = AssessmentType.query.get_or_404(assessment_type_id)

    if request.method == 'GET':  # Handle GET request (display the form)
        form = AssessmentTypeForm() # instantiate form
        form.name.data = assessment_type.name  # Pre-populate the form
        return render_template('edit_assessment_type.html', assessment_type=assessment_type, form=form)

    if request.method == 'POST':
        form = AssessmentTypeForm(request.form) # Get data from the form
        if form.validate_on_submit(): # Validate the form
            assessment_type.name = form.name.data # Update the name
            try:
                db.session.commit()
                flash('Assessment type updated successfully!', 'success')
                return redirect(url_for('main.assessment_types'))
            except Exception as e:
                db.session.rollback()
                flash(f'An error occurred: {str(e)}', 'danger')
        else:
            flash('Error in form', 'danger') #Flash error if the form is not valid
            return render_template('edit_assessment_type.html', assessment_type=assessment_type, form=form) #Render the form again with error


@main.route('/delete_assessment_type/<int:assessment_type_id>', methods=['GET', 'POST'])
@login_required
def delete_assessment_type(assessment_type_id):
    if not current_user.role == 'admin':
        flash('Access denied!', 'danger')
        return redirect(url_for('main.index'))

    assessment_type = AssessmentType.query.get_or_404(assessment_type_id)
    try:
        db.session.delete(assessment_type)
        db.session.commit()
        flash('Assessment type deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'An error occurred: {str(e)}', 'danger')

    return redirect(url_for('main.assessment_types'))

@main.route('/assessments')
@login_required
def assessments():
    all_assessments = Assessment.query.join(Class, Assessment.class_id == Class.id) \
        .join(AssessmentType, Assessment.assessment_type_id == AssessmentType.id) \
        .filter(Class.school_id == current_user.school_id).add_columns(
            Assessment.id,
            Assessment.name,
            Assessment.date,
            Assessment.academic_session,
            Assessment.term,
            Class.class_name.label('class_name'),
            AssessmentType.name.label('assessment_type_name')
        ) \
        .all()

    return render_template('assessments.html', assessments=all_assessments)

@main.route('/add_assessment', methods=['GET', 'POST'])
@login_required
def add_assessment():
    form = AssessmentForm()
    form.assessment_type.choices = [(atype.id, atype.name) for atype in AssessmentType.query.all()]
    form.class_id.choices = [(cls.id, cls.class_name) for cls in Class.query.all()]

    if form.validate_on_submit():
        new_assessment = Assessment(
            name=form.name.data,
            date=form.date.data,
            assessment_type_id=form.assessment_type.data,
            class_id=form.class_id.data,
            academic_session=form.academic_session.data,
            term=form.term.data
        )
        db.session.add(new_assessment)
        db.session.commit()
        flash('Assessment added successfully!', 'success')
        return redirect(url_for('main.assessments'))

    return render_template('add_assessment.html', form=form)


@main.route('/edit_assessment/<int:assessment_id>', methods=['GET', 'POST'])
@login_required
def edit_assessment(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    form = AssessmentForm(obj=assessment)
    form.assessment_type.choices = [(atype.id, atype.name) for atype in AssessmentType.query.all()]
    form.class_id.choices = [(cls.id, cls.class_name) for cls in Class.query.all()]

    if form.validate_on_submit():
        assessment.name = form.name.data
        assessment.date = form.date.data
        assessment.assessment_type_id = form.assessment_type.data
        assessment.class_id = form.class_id.data
        assessment.academic_session = form.academic_session.data
        assessment.term = form.term.data

        db.session.commit()
        flash('Assessment updated successfully!', 'success')
        return redirect(url_for('main.assessments'))

    return render_template('edit_assessment.html', form=form, assessment=assessment)

@main.route('/delete_assessment/<int:assessment_id>', methods=['POST'])
@login_required
def delete_assessment(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    db.session.delete(assessment)
    db.session.commit()
    flash('Assessment deleted successfully!', 'success')
    return redirect(url_for('main.assessments'))


# @main.route('/assessment_subject_scores')
# def assessment_subject_scores():
#     # Get filter parameters from the request
#     assessment_id = request.args.get('assessment')
#     academic_session = request.args.get('academic_session')
#     subject_id = request.args.get('subject')
#     student_id = request.args.get('student')
    

#     # Base query
#     query = db.session.query(
#         AssessmentSubjectScore,
#         Assessment.name.label('assessment_name'),
#         Subject.name.label('subject_name'),
#         Student.first_name,
#         Student.last_name,
#         AssessmentSubjectScore.total_marks,
#         Assessment.academic_session
#     ).join(
#         Assessment, AssessmentSubjectScore.assessment_id == Assessment.id
#     ).join(
#         Subject, AssessmentSubjectScore.subject_id == Subject.id
#     ).join(
#         Student, AssessmentSubjectScore.student_id == Student.id
#     )

#     # Apply filters
#     if assessment_id:
#         query = query.filter(AssessmentSubjectScore.assessment_id == assessment_id)
#     if academic_session:
#         query = query.filter(Assessment.academic_session == academic_session)
#     if subject_id:
#         query = query.filter(AssessmentSubjectScore.subject_id == subject_id)
#     if student_id:
#         query = query.filter(AssessmentSubjectScore.student_id == student_id)

#     # Fetch filtered results
#     scores = query.all()

#     # Fetch filter options
#     assessments = Assessment.query.all()
#     academic_sessions = db.session.query(Assessment.academic_session).distinct().all()
#     subjects = Subject.query.all()
#     students = Student.query.all()

#     return render_template(
#         'assessment_subject_scores.html',
#         scores=scores,
#         assessments=assessments,
#         academic_sessions=[session[0] for session in academic_sessions],
#         subjects=subjects,
#         students=students
#     )

@main.route('/assessment_subject_scores')
@login_required
def assessment_subject_scores():
    # Ensure the user is linked to a school
    if not current_user.school_id:
        flash("No school linked to your account.", "danger")
        return redirect(url_for('main.index'))

    school_id = current_user.school_id  # Get logged-in user's school

    # Get filter parameters from the request
    assessment_id = request.args.get('assessment')
    academic_session = request.args.get('academic_session')
    subject_id = request.args.get('subject')
    student_id = request.args.get('student')

    # Base query with school filtering
    query = db.session.query(
        AssessmentSubjectScore,
        Assessment.name.label('assessment_name'),
        Subject.name.label('subject_name'),
        Student.first_name,
        Student.last_name,
        AssessmentSubjectScore.total_marks,
        Assessment.academic_session
    ).join(
        Assessment, AssessmentSubjectScore.assessment_id == Assessment.id
    ).join(
        Subject, AssessmentSubjectScore.subject_id == Subject.id
    ).join(
        Student, AssessmentSubjectScore.student_id == Student.id
    ).filter(
        Student.school_id == school_id  # Ensuring students are from the logged-in user's school
    )

    # Apply additional filters
    if assessment_id:
        query = query.filter(AssessmentSubjectScore.assessment_id == assessment_id)
    if academic_session:
        query = query.filter(Assessment.academic_session == academic_session)
    if subject_id:
        query = query.filter(AssessmentSubjectScore.subject_id == subject_id)
    if student_id:
        query = query.filter(AssessmentSubjectScore.student_id == student_id)

    # Fetch filtered results
    scores = query.all()

    # Fetch filter options, ensuring they are linked to the same school
    assessments = Assessment.query.filter_by(school_id=school_id).all()
    academic_sessions = db.session.query(Assessment.academic_session).filter_by(school_id=school_id).distinct().all()
    subjects = Subject.query.filter_by(school_id=school_id).all()
    students = Student.query.filter_by(school_id=school_id).all()

    return render_template(
        'assessment_subject_scores.html',
        scores=scores,
        assessments=assessments,
        academic_sessions=[session[0] for session in academic_sessions],
        subjects=subjects,
        students=students
    )


# @main.route('/add_assessment_score', methods=['GET', 'POST'])
# def add_assessment_score():
#     if request.method == 'POST':
#         try:
#             # Fetch the data from the form
#             academic_session = request.form.get('academic_session')
#             assessment_id = int(request.form.get('assessment_id'))
#             student_id = int(request.form.get('student_id'))
#             class_id = int(request.form.get('class_id'))

#             # Get the list of selected subjects (only subjects with checked checkboxes)
#             selected_subject_ids = request.form.getlist('selected_subjects')

#             # Loop through the selected subjects and collect scores
#             for subject_id in selected_subject_ids:
#                 score_key = f"score_{subject_id}"
#                 score_value = request.form.get(score_key)

#                 if score_value:
#                     try:
#                         total_marks = int(score_value)
#                         # Add score entry to the database
#                         new_score = AssessmentSubjectScore(
#                             assessment_id=assessment_id,
#                             subject_id=int(subject_id),
#                             student_id=student_id,
#                             total_marks=total_marks
#                         )
#                         db.session.add(new_score)
#                     except ValueError:
#                         continue

#             db.session.commit()
#             flash("Assessment scores added successfully!", "success")
#             return redirect(url_for('main.assessment_subject_scores'))
#         except Exception as e:
#             db.session.rollback()
#             flash(f"Error adding assessment scores: {e}", "danger")
#             return redirect(url_for('main.add_assessment_score'))

#     # Fetch academic sessions, assessments, subjects, and students
#     academic_sessions = list(set(assessment.academic_session for assessment in Assessment.query.all()))
#     assessments = Assessment.query.all()
#     subjects = Subject.query.all()
#     students = Student.query.all()
#     classes = Class.query.all()

#     return render_template(
#         'add_assessment_subject_score.html',
#         academic_sessions=academic_sessions,
#         assessments=assessments,
#         subjects=subjects,
#         students=students,
#         classes=classes
#     )

@main.route('/add_assessment_score', methods=['GET', 'POST'])
@login_required
def add_assessment_score():
    if not current_user.school_id:
        flash("You are not linked to any school.", "danger")
        return redirect(url_for('main.index'))

    school_id = current_user.school_id  # Get the logged-in user's school ID

    if request.method == 'POST':
        try:
            # Fetch the data from the form
            academic_session = request.form.get('academic_session')
            assessment_id = int(request.form.get('assessment_id'))
            student_id = int(request.form.get('student_id'))
            class_id = int(request.form.get('class_id'))

            # Get the list of selected subjects (only subjects with checked checkboxes)
            selected_subject_ids = request.form.getlist('selected_subjects')

            # Ensure the assessment and student belong to the logged-in user's school
            assessment = Assessment.query.filter_by(id=assessment_id, school_id=school_id).first()
            student = Student.query.filter_by(id=student_id, school_id=school_id).first()

            if not assessment or not student:
                flash("Invalid selection. Please choose valid records from your school.", "danger")
                return redirect(url_for('main.add_assessment_score'))

            # Loop through the selected subjects and collect scores
            for subject_id in selected_subject_ids:
                score_key = f"score_{subject_id}"
                score_value = request.form.get(score_key)

                if score_value:
                    try:
                        total_marks = int(score_value)

                        # Ensure the subject belongs to the school
                        subject = Subject.query.filter_by(id=int(subject_id), school_id=school_id).first()
                        if not subject:
                            flash(f"Invalid subject selection: {subject_id}.", "danger")
                            continue

                        # Add score entry to the database
                        new_score = AssessmentSubjectScore(
                            assessment_id=assessment_id,
                            subject_id=int(subject_id),
                            student_id=student_id,
                            total_marks=total_marks
                        )
                        db.session.add(new_score)

                    except ValueError:
                        flash(f"Invalid score for subject ID {subject_id}.", "danger")
                        continue

            db.session.commit()
            flash("Assessment scores added successfully!", "success")
            return redirect(url_for('main.assessment_subject_scores'))

        except Exception as e:
            db.session.rollback()
            flash(f"Error adding assessment scores: {e}", "danger")
            return redirect(url_for('main.add_assessment_score'))

    # Fetch only data belonging to the logged-in user's school
    academic_sessions = list(set(assessment.academic_session for assessment in Assessment.query.filter_by(school_id=school_id).all()))
    assessments = Assessment.query.filter_by(school_id=school_id).all()
    subjects = Subject.query.filter_by(school_id=school_id).all()
    students = Student.query.filter_by(school_id=school_id).all()
    classes = Class.query.filter_by(school_id=school_id).all()

    return render_template(
        'add_assessment_subject_score.html',
        academic_sessions=academic_sessions,
        assessments=assessments,
        subjects=subjects,
        students=students,
        classes=classes
    )

@main.route('/assessment_subject_scores/edit/<int:id>', methods=['GET', 'POST'])
def edit_assessment_subject_score(id):
    # Fetch the record to be edited
    score = AssessmentSubjectScore.query.get_or_404(id)

    if request.method == 'POST':
        # Update the score's fields from the form data
        score.subject_id = request.form['subject_id']
        score.total_marks = request.form['total_marks']

        try:
            db.session.commit()
            flash('Assessment Subject Score updated successfully!', 'success')
            return redirect(url_for('main.assessment_subject_scores'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating score: {str(e)}', 'danger')

    # Fetch data for dropdowns
    subjects = Subject.query.all()
    assessments = Assessment.query.all()

    return render_template(
        'edit_assessment_subject_score.html',
        score=score,
        subjects=subjects,
        assessments=assessments
    )


@main.route('/assessment_subject_scores/delete/<int:id>', methods=['POST'])
def delete_assessment_subject_score(id):
    score = AssessmentSubjectScore.query.get_or_404(id)

    try:
        db.session.refresh(score)  # Refresh the object from the database
        db.session.delete(score)
        db.session.commit()
        flash('Assessment Subject Score deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting score: {str(e)}', 'danger')

    return redirect(url_for('main.assessment_subject_scores'))


@main.route('/assessment_subject_scores/delete', methods=['POST'])
def delete_multiple_assessment_subject_scores():
    selected_ids = request.form.getlist('delete_ids')
    
    if not selected_ids:
        flash('No assessment scores selected for deletion.', 'warning')
        return redirect(url_for('main.assessment_subject_scores'))
    
    try:
        for score_id in selected_ids:
            score = AssessmentSubjectScore.query.get(score_id)
            if score:
                db.session.delete(score)
        
        db.session.commit()
        flash(f'Successfully deleted {len(selected_ids)} assessment score(s).', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting scores: {str(e)}', 'danger')
    
    return redirect(url_for('main.assessment_subject_scores'))



# Route to view all subjects for a class
@main.route('/class_subjects/<int:class_id>')
@login_required
def class_subjects(class_id):
    class_instance = Class.query.get_or_404(class_id)
    subjects = ClassSubject.query.filter_by(class_id=class_id).all()
    return render_template('class_subjects.html', class_instance=class_instance, subjects=subjects)

# Route to add a subject to a class
@main.route('/add_class_subject/<int:class_id>', methods=['GET', 'POST'])
@login_required
def add_class_subject(class_id):
    if not current_user.role == 'admin':
        flash('Access denied!', 'danger')
        return redirect(url_for('main.index'))
    class_instance = Class.query.get_or_404(class_id)
    if request.method == 'POST':
        subject_id = request.form.get('subject_id')
        if not subject_id:
            flash('Subject is required.', 'danger')
        else:
            new_class_subject = ClassSubject(
                class_id=class_id,
                subject_id=subject_id
            )
            db.session.add(new_class_subject)
            db.session.commit()
            flash('Subject added to class successfully!', 'success')
            return redirect(url_for('main.class_subjects', class_id=class_id))
    subjects = Subject.query.filter_by(school_id=current_user.school_id).all()
    return render_template('add_class_subject.html', class_instance=class_instance, subjects=subjects)



def calculate_grade(total_score):
    if 70 <= total_score <= 100:
        return 'A', 'Excellent'
    elif 60 <= total_score < 70:
        return 'B', 'Very Good'
    elif 50 <= total_score < 60:
        return 'C', 'Good'
    elif 45 <= total_score < 50:
        return 'D', 'Pass'
    else:
        return 'E', 'Weak'


def create_student_result_pdf(student):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    styles = getSampleStyleSheet()

    # Title Section (with Logo and Colored School Name)
    pdf.setFont("Helvetica-Bold", 14)

    # School Logo (Top Left)
    try:
        logo_path = student.get("school_logo")  # Access from dictionary
        if logo_path:
            full_logo_path = os.path.join(Config.UPLOAD_FOLDER, logo_path)
            if os.path.exists(full_logo_path):
                img = ImageReader(full_logo_path)
                pdf.drawImage(img, 50, 750, width=75, height=75, preserveAspectRatio=True)
            else:
                print(f"Logo file not found: {full_logo_path}")
                pdf.drawString(50, 775, "School Logo Not Found")
        else:
            print("No logo path provided.")
            pdf.drawString(50, 775, "No School Logo")
    except Exception as e:
        print(f"Error loading logo: {e}")
        pdf.drawString(50, 775, "Error Loading Logo")

    # School Name (Colored)
    pdf.setFillColor(colors.red)
    pdf.drawCentredString(width / 2, 800, student['school_name'])
    pdf.setFillColor(colors.black)

    pdf.setFont("Helvetica", 12)
    pdf.drawCentredString(width / 2, 780, student['school_address'])
    pdf.drawCentredString(width / 2, 760, "INDIVIDUAL STUDENT'S ASSESSMENT SHEET")
    pdf.drawCentredString(width / 2, 740, "JUNIOR SECONDARY SCHOOL")

    student_details_data = [
        ["RESULT ID:", student['result_id'], "SEX:", student['sex'], "TERM:", student['term']],
        ["NAME OF STUDENT:", student['name'], "SESSION:", student['session'], "CLASS:", student['class_name']],
        [
            Paragraph("NO. OF TIMES<br/>PRESENT", styles['Normal']),
            student['times_present'],
            Paragraph("NO. OF TIMES<br/>SCHOOL OPENED", styles['Normal']),
            student['times_opened'], "POSITION:", student['class_name']
        ],
    ]

    col_widths = [120, 130, 60, 60, 80, 80]
    student_details_table = Table(student_details_data, colWidths=col_widths)

    student_details_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
    ]))

    student_details_table.wrapOn(pdf, width, height)
    student_details_table.drawOn(pdf, 30, 630)

    subject_data = []
    grand_total = 0
    passes = 0

    for subject in student['subjects']:
        total_score = subject.get('CA1', 0) + subject.get('CA2', 0) + subject.get('EXAMINATION', 0)
        grade, remark = calculate_grade(total_score)
        grand_total += total_score
        if grade in ['A', 'B', 'C']:
            passes += 1
        subject_data.append([subject['name'], subject.get('CA1', 0), subject.get('CA2', 0), subject.get('EXAMINATION', 0), total_score, grade, remark])

    overall_grade = calculate_grade(grand_total / len(student['subjects']))
    student['overall_grade'] = overall_grade[0]

    header_row_1 = [
        "SUBJECT",
        Paragraph("CONTINUOUS ASSESSMENT", styles["h3"]),
        Spacer(1, 1), Spacer(1, 1), Spacer(1, 1),
        Paragraph("TERM SUMMARY", styles["h3"]),
        Spacer(1, 1),
    ]

    header_row_2 = ["SUBJECT", "CA1", "CA2", "EXAMINATION", "TOTAL", "GRADE", "REMARK"]
    full_subject_data = [header_row_1, header_row_2] + subject_data

    row_heights = [20, 20] + [20] * len(subject_data)
    subject_table = Table(full_subject_data, colWidths=[150, 50, 50, 80, 50, 50, 100], rowHeights=row_heights)

    subject_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (1, 2), (4, -1), colors.beige),
        ('BACKGROUND', (5, 2), (6, -1), colors.beige),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (1, 0), (4, 0), colors.lightgrey),
        ('BACKGROUND', (5, 0), (6, 0), colors.lightgrey),
        ('SPAN', (1, 0), (4, 0)),
        ('SPAN', (5, 0), (6, 0)),
        ('BACKGROUND', (0, 1), (-1, 1), colors.lightgrey),
    ]))

    student_details_table.wrapOn(pdf, width, height)
    student_details_table_height = student_details_table._height
    subject_table_y_position = 390 - student_details_table_height - 20

    subject_table.wrapOn(pdf, width, height)
    subject_table.drawOn(pdf, 30, subject_table_y_position)

    term_average = grand_total / len(student['subjects']) if student['subjects'] else 0

    summary_data = [
        ["Term Grand Total:", str(grand_total)],
        ["Term Average:", f"{term_average:.2f}"],
        ["Number of Subject Passes:", str(passes)],
        ["Overall Grade:", student.get("overall_grade", "N/A")],
    ]

    grading_key_data = [
        ["Figures", "Grade", "Remark"],
        ["70 - 100", "A", "Excellent"],
        ["60 - 69", "B", "Very Good"],
        ["50 - 59", "C", "Good"],
        ["45 - 49", "D", "Pass"],
        ["0 - 44", "E", "Weak"],
    ]

    summary_table = Table(summary_data, colWidths=[150, 80])
    grading_key_table = Table(grading_key_data, colWidths=[80, 60, 80])

    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))

    grading_key_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
    ]))

    summary_table.wrapOn(pdf, width, height)
    summary_table.drawOn(pdf, 30, 180)

    grading_key_table.wrapOn(pdf, width, height)
    grading_key_table.drawOn(pdf, 300, 160)

    pdf.setFont("Helvetica", 10)
    pdf.drawString(50, 120, "Class Teacher Remark: __________________________")
    pdf.drawString(50, 100, "Principal Remark: __________________________")

    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawCentredString(width / 2, 62, "____________________________")
    pdf.drawCentredString(width / 2, 50, "PRINCIPAL'S SIGNATURE")

    pdf.save()
    buffer.seek(0)
    return buffer




@main.route('/download_result_sheet/<int:student_id>')
def download_result_sheet(student_id):
    try:
        student_record = Student.query.filter_by(id=student_id).first()
        if not student_record:
            return jsonify({"error": "Student not found"}), 404

        school_record = School.query.filter_by(id=student_record.school_id).first()
        if not school_record:
            return jsonify({"error": "School not found"}), 404

        subjects = AssessmentSubjectScore.query.filter_by(student_id=student_id).all()
        print(f"Subjects for student {student_id}: {subjects}")  # Debug log

        assessments = Assessment.query.filter_by(class_id=student_record.class_id).all()
        print(f"Assessments for class {student_record.class_id}: {assessments}")  # Debug log

        attendance_records = Attendance.query.filter_by(student_id=student_id).all()
        print(f"Attendance Records for student {student_id}: {attendance_records}")  # Debug log

        times_present = sum(1 for record in attendance_records if record.days_present == 'present')
        times_opened = len(attendance_records)

        subject_scores = {}
        for subject_score in subjects:
            subject_name = subject_score.subject.name
            if subject_name not in subject_scores:
                subject_scores[subject_name] = {'CA1': 0, 'CA2': 0, 'EXAMINATION': 0}

            for assessment in assessments:
                if assessment.id == subject_score.assessment_id:
                    print(f"Assessment Name: {assessment.name}, Marks: {subject_score.total_marks}")  # Debug log
                    if 'CA1' in assessment.name:
                        subject_scores[subject_name]['CA1'] = subject_score.total_marks
                    elif 'CA2' in assessment.name:
                        subject_scores[subject_name]['CA2'] = subject_score.total_marks
                    elif 'EXAMINATION' in assessment.name:
                        subject_scores[subject_name]['EXAMINATION'] = subject_score.total_marks

        print(f"Subject Scores: {subject_scores}")  # Debug log

        grand_total = sum(sum(scores.values()) for scores in subject_scores.values())
        print(f"Grand Total: {grand_total}")  # Debug log

        student = {
            "result_id": student_record.id,
            "name": f"{student_record.first_name} {student_record.last_name}",
            "sex": student_record.gender,
            "term": assessments[0].term if assessments else "N/A",
            "session": assessments[0].academic_session if assessments else "N/A",
            "times_present": times_present,
            "times_opened": times_opened,
            "class_name": student_record.class_.class_name,
            "subjects": [
                {
                    "name": subject,
                    "CA1": scores['CA1'],
                    "CA2": scores['CA2'],
                    "EXAMINATION": scores['EXAMINATION']
                }
                for subject, scores in subject_scores.items()
            ],
            "grand_total": grand_total,
            "overall_grade": "A",
            "next_term": "2023-09-01",
            "school_name": school_record.name,
            "school_address": school_record.address,
            "school_logo": school_record.school_logo  # Add this line
        }

        pdf_buffer = create_student_result_pdf(student)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=f"{student['name']}_result.pdf",
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# @main.route('/assign_teachers', methods=['GET', 'POST'])
# @login_required
# def assign_teachers():
#     # Create the form
#     form = AssignTeachersForm()

#     # Get the classes and teachers for the current school
#     classes = Class.query.filter_by(school_id=current_user.school_id).all()
#     teachers = Teacher.query.filter_by(school_id=current_user.school_id).all()

#     # Populate the form choices dynamically
#     form.class_id.choices = [(cls.id, cls.class_name) for cls in classes]
#     form.teacher_ids.choices = [(teacher.id, f"{teacher.first_name} {teacher.last_name}") for teacher in teachers]

#     if form.validate_on_submit():
#         class_id = form.class_id.data
#         selected_teacher_ids = form.teacher_ids.data

#         # Fetch the class and selected teachers
#         selected_class = Class.query.get(class_id)
#         selected_teachers = Teacher.query.filter(Teacher.id.in_(selected_teacher_ids)).all()

#         # Clear existing teachers to avoid duplicates
#         selected_class.teachers = []

#         # Add new teachers to the class
#         selected_class.teachers.extend(selected_teachers)

#         try:
#             db.session.commit()
#             flash('Teachers assigned successfully!', 'success')
#             return redirect(url_for('main.classes'))
#         except Exception as e:
#             db.session.rollback()
#             flash(f'Error assigning teachers: {str(e)}', 'danger')


#     return render_template('assign_teachers.html', form=form)


@main.route('/assign_teachers', methods=['GET', 'POST'])
@login_required
def assign_teachers():
    form = AssignTeachersForm()

    current_school = current_user.school
    if not current_school:
        flash("You are not associated with any school.", "warning")
        return redirect(url_for('main.index'))  # Or appropriate route

    # Filter classes and teachers by the logged-in user's school
    classes = Class.query.filter_by(school_id=current_school.id).all()
    teachers = Teacher.query.filter_by(school_id=current_school.id).all()

    form.class_id.choices = [(cls.id, cls.class_name) for cls in classes]
    form.teacher_ids.choices = [(teacher.id, f"{teacher.first_name} {teacher.last_name}") for teacher in teachers]

    if form.validate_on_submit():
        class_id = form.class_id.data
        selected_teacher_ids = form.teacher_ids.data

        selected_class = Class.query.get_or_404(class_id) # Add 404 handling

        if selected_class.school_id != current_school.id: # Check school association
             flash("You do not have permission to modify this class.", "danger")
             return redirect(url_for('main.classes'))

        selected_teachers = Teacher.query.filter(Teacher.id.in_(selected_teacher_ids)).all()

        # Check if all selected teachers belong to the current school
        for teacher in selected_teachers:
            if teacher.school_id != current_school.id:
                flash("One or more selected teachers do not belong to your school.", "danger")
                return redirect(url_for('main.assign_teachers'))  # Redirect back to the form

        selected_class.teachers = selected_teachers  # Simplified assignment

        try:
            db.session.commit()
            flash('Teachers assigned successfully!', 'success')
            return redirect(url_for('main.classes'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error assigning teachers: {str(e)}', 'danger')

    return render_template('assign_teachers.html', form=form)



# School Fees Management

fee_components = {
    "Tuition Fee": "Fee for academic instruction and resources.",
    "Registration Fee": "Fee for enrollment in the academic year.",
    "PTA Levy": "Parent-Teacher Association levy for school activities.",
    "Sports Fee": "Fee for participation in sports and physical activities.",
    "Uniform Fee": "Cost of school uniforms.",
    "Textbooks/Workbooks": "Cost of textbooks and workbooks for the academic year.",
    "Examination Fee": "Fee for conducting examinations and assessments.",
    "ICT Fee": "Fee for access to ICT facilities and resources.",
    "Medical Fee": "Fee for access to medical facilities and services.",
    "Miscellaneous": "Additional fees for other school-related activities."
}

# Save as JSON file
with open("fee_components.json", "w") as file:
    json.dump(fee_components, file, indent=4)


# @main.route('/fee_components')
# def fee_components():
#     # Fetch all fee components from the database
#     components = FeeComponent.query.all()
#     return render_template('fee_components.html', components=components)

@main.route('/fee_components')
@login_required
def fee_components():
    # Ensure the user is linked to a school
    if not current_user.school_id:
        flash("No school linked to your account.", "danger")
        return redirect(url_for('main.index'))

    school_id = current_user.school_id  # Get logged-in user's school

    # Fetch only fee components related to the user's school
    components = FeeComponent.query.filter_by(school_id=school_id).all()

    return render_template('fee_components.html', components=components)


# Load fee components from the JSON file
with open("fee_components.json", "r") as file:
    fee_components_from_json = json.load(file)


# @main.route('/add_fee_component', methods=['GET', 'POST'])
# def add_fee_component():
#     if request.method == 'POST':
#         # Get form data
#         name = request.form['name']
#         description = request.form['description']
#         school_id = request.form['school_id']

#         # Create a new FeeComponent
#         component = FeeComponent(
#             name=name,
#             description=description,
#             school_id=school_id
#         )
#         db.session.add(component)
#         db.session.commit()
#         flash('Fee component added successfully!', 'success')
#         return redirect(url_for('main.fee_components'))

#     # Fetch all schools for the dropdown
#     schools = School.query.all()
#     return render_template('add_fee_component.html', schools=schools)



# @main.route('/add_fee_component_to_class', methods=['GET', 'POST'])
# def add_fee_component_to_class():
#     if request.method == 'POST':
#         # Get the list of selected class IDs
#         selected_classes = request.form.getlist('selected_classes')

#         # Get the list of selected component IDs
#         selected_components = request.form.getlist('selected_components')

#         # Loop through all selected classes and fee components
#         for class_id in selected_classes:
#             for component_id in selected_components:
#                 # Get the amount for the selected component
#                 amount_field = f"amount_{component_id}"
#                 if amount_field in request.form and request.form[amount_field]:
#                     amount = float(request.form[amount_field])

#                     # Add the fee component to the database for the current class
#                     class_fee_component = ClassFeeComponent(
#                         class_id=int(class_id),
#                         component_id=int(component_id),
#                         amount=amount
#                     )
#                     db.session.add(class_fee_component)

#         # Commit all changes to the database
#         db.session.commit()
#         flash('Selected fee components added to the selected classes successfully!', 'success')
#         return redirect(url_for('main.view_class_fees', class_id=class_id))  # Adjust the redirect as needed
    
#     # Fetch all classes and fee components for the form
#     classes = Class.query.all()
#     components = FeeComponent.query.all()
#     return render_template('add_fee_component_to_class.html', classes=classes, components=components)


# @main.route('/edit_fee_component/<int:id>', methods=['GET', 'POST'])
# def edit_fee_component(id):
#     component = FeeComponent.query.get_or_404(id)
#     if request.method == 'POST':
#         # Update the fee component
#         component.name = request.form['name']
#         component.description = request.form['description']
#         component.school_id = request.form['school_id']
#         db.session.commit()
#         flash('Fee component updated successfully!', 'success')
#         return redirect(url_for('main.fee_components'))

#     # Fetch all schools for the dropdown
#     schools = School.query.all()
#     return render_template('edit_fee_component.html', component=component, schools=schools)

# @main.route('/delete_fee_component/<int:id>', methods=['POST'])
# def delete_fee_component(id):
#     component = FeeComponent.query.get_or_404(id)
#     db.session.delete(component)
#     db.session.commit()
#     flash('Fee component deleted successfully!', 'success')
#     return redirect(url_for('main.fee_components'))

@main.route('/add_fee_component', methods=['GET', 'POST'])
@login_required  # Protect the route
def add_fee_component():
    if request.method == 'POST':
        name = request.form['name']
        description = request.form['description']
        academic_year = request.form['academic_year']
        term = request.form['term']

        # Use current_user to get the logged-in user's school
        school_id = current_user.school_id  # Assuming your User model has a school_id

        component = FeeComponent(
            name=name,
            description=description,
            school_id=school_id,
            academic_year=academic_year,
            term=term
        )
        db.session.add(component)
        db.session.commit()
        flash('Fee component added successfully!', 'success')
        return redirect(url_for('main.fee_components'))

    # Filter schools based on the logged-in user
    schools = School.query.filter_by(id=current_user.school_id).all() # Only show the logged in user's school
    return render_template('add_fee_component.html', schools=schools)



# @main.route('/add_fee_component_to_class', methods=['GET', 'POST'])
# @login_required
# def add_fee_component_to_class():
#     if request.method == 'POST':
#         selected_classes = request.form.getlist('selected_classes')
#         selected_components = request.form.getlist('selected_components')

#         for class_id in selected_classes:
#             for component_id in selected_components:
#                 amount_field = f"amount_{component_id}"
#                 if amount_field in request.form and request.form[amount_field]:
#                     amount = float(request.form[amount_field])

#                     # Verify that the class belongs to the user's school
#                     class_obj = Class.query.get(class_id)
#                     if class_obj and class_obj.school_id == current_user.school_id: #Added check to ensure the class belongs to the logged in user's school
#                         class_fee_component = ClassFeeComponent(
#                             class_id=int(class_id),
#                             component_id=int(component_id),
#                             amount=amount
#                         )
#                         db.session.add(class_fee_component)
#                     else:
#                         flash(f"Error: Class with ID {class_id} does not belong to your school.", 'danger')  # Or handle the error as you see fit.
#                         return redirect(url_for('main.add_fee_component_to_class')) # Redirect back to the form

#         db.session.commit()
#         flash('Selected fee components added to the selected classes successfully!', 'success')
#         # You'll need to determine the correct redirect.  Perhaps a view for a specific class's fees?
#         return redirect(url_for('main.fee_components'))  # Or a more appropriate redirect

#     # Filter classes and components by school
#     classes = Class.query.filter_by(school_id=current_user.school_id).all()
#     components = FeeComponent.query.filter_by(school_id=current_user.school_id).all()
#     return render_template('add_fee_component_to_class.html', classes=classes, components=components)

# @main.route('/add_fee_component_to_class', methods=['GET', 'POST']) # This route will handle the POST request
# @login_required
# def add_fee_component_to_class():
#     if request.method == 'POST':
#         selected_classes = request.form.getlist('selected_classes')
#         selected_components = request.form.getlist('selected_components')
#         academic_year = request.form['academic_year']
#         term = request.form['term']


#         # make it possible for selection of academic year

#         for class_id in selected_classes:
#             for component_id in selected_components:
#                 amount_field = f"amount_{component_id}"
#                 if amount_field in request.form and request.form[amount_field]:
#                     amount = float(request.form[amount_field])

#                     class_obj = Class.query.get(class_id)
#                     if class_obj and class_obj.school_id == current_user.school_id:
#                         component = FeeComponent.query.get(component_id)
#                         if component:
#                             # Check if a fee component with the same name, academic year, and term already exists for the class
#                             existing_fee = ClassFeeComponent.query.join(FeeComponent).filter(
#                                 ClassFeeComponent.class_id == class_id,
#                                 FeeComponent.name == component.name,
#                                 FeeComponent.academic_year == academic_year,
#                                 FeeComponent.term == term
#                             ).first()

#                             if existing_fee:
#                                 # Update the existing fee
#                                 existing_fee.amount = amount
#                             else:
#                                 # Create a new fee component
#                                 fee_component = FeeComponent(
#                                     name=component.name,
#                                     description=component.description,
#                                     school_id=current_user.school_id,
#                                     academic_year=academic_year,
#                                     term=term
#                                 )
#                                 db.session.add(fee_component)
#                                 db.session.flush() # Get the ID for the new fee_component
#                                 class_fee_component = ClassFeeComponent(
#                                     class_id=int(class_id),
#                                     component_id=fee_component.id, # Use fee_component.id
#                                     amount=amount
#                                 )
#                                 db.session.add(class_fee_component)
#                         else:
#                             flash(f"Error: Fee component with ID {component_id} not found.", 'danger')
#                             return redirect(url_for('main.add_fee_component_to_class'))
#                     else:
#                         flash(f"Error: Class with ID {class_id} does not belong to your school.", 'danger')
#                         return redirect(url_for('main.add_fee_component_to_class'))

#         db.session.commit()
#         flash('Selected fee components added to the selected classes successfully!', 'success')
#         return redirect(url_for('main.generate_fees'))

#     academic_years = [
#         year.academic_year
#         for year in db.session.query(FeeComponent.academic_year).filter_by(school_id=current_user.school_id).distinct().all()
#     ]

#     # Fetch all classes and fee components for the form
#     classes = Class.query.filter_by(school_id=current_user.school_id).all()
#     components = FeeComponent.query.filter_by(school_id=current_user.school_id).all()
#     return render_template('add_fee_component_to_class.html', classes=classes, components=components, academic_years=academic_years)


@main.route('/add_fee_component_to_class', methods=['GET', 'POST'])
@login_required
def add_fee_component_to_class():
    school = School.query.filter_by(id=current_user.school_id).first()

    if not school:
        flash("No school assigned to your account.", "danger")
        return redirect(url_for('main.index'))

    if request.method == 'POST':
        selected_classes = request.form.getlist('selected_classes')
        selected_components = request.form.getlist('selected_components')
        academic_year = request.form['academic_year']
        term = request.form['term']

        for class_id in selected_classes:
            for component_id in selected_components:
                amount_field = f"amount_{component_id}"
                if amount_field in request.form and request.form[amount_field]:
                    amount = float(request.form[amount_field])

                    class_obj = Class.query.get(class_id)
                    if class_obj and class_obj.school_id == current_user.school_id:
                        component = FeeComponent.query.get(component_id)
                        if component and component.school_id == current_user.school_id:
                            # Check if a ClassFeeComponent entry *already exists* for this class, component, year, and term
                            existing_class_fee = ClassFeeComponent.query.join(FeeComponent).filter(
                                ClassFeeComponent.class_id == class_id,
                                FeeComponent.id == component_id,  # Filter by component ID
                                FeeComponent.academic_year == academic_year,
                                FeeComponent.term == term,
                                FeeComponent.school_id == current_user.school_id
                            ).first()

                            if existing_class_fee:
                                # Update the existing ClassFeeComponent
                                existing_class_fee.amount = amount
                            else:
                                # Create a new ClassFeeComponent ONLY (no new FeeComponent)
                                class_fee_component = ClassFeeComponent(
                                    class_id=int(class_id),
                                    component_id=int(component_id),  # Use the existing component ID
                                    amount=amount
                                )
                                db.session.add(class_fee_component)
                        else:
                            flash(f"Error: Fee component with ID {component_id} not found or doesn't belong to your school.", 'danger')
                            return redirect(url_for('main.add_fee_component_to_class'))
                    else:
                        flash(f"Error: Class with ID {class_id} does not belong to your school.", 'danger')
                        return redirect(url_for('main.add_fee_component_to_class'))

        db.session.commit()
        flash('Fee components added/updated successfully!', 'success')  # More accurate message
        return redirect(url_for('main.generate_fees'))

    # GET request handling (same as before)
    academic_years = [
        year.academic_year
        for year in db.session.query(FeeComponent.academic_year).filter_by(school_id=current_user.school_id).distinct().all()
    ]


    classes = Class.query.filter_by(school_id=school.id).all()
    components = FeeComponent.query.filter_by(school_id=school.id).all()
    return render_template('add_fee_component_to_class.html', classes=classes, components=components, academic_years=academic_years)

@main.route('/edit_fee_component/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_fee_component(id):
    component = FeeComponent.query.get_or_404(id)

    # Make sure the component belongs to the logged-in user's school
    if component.school_id != current_user.school_id:
        flash("You don't have permission to edit this fee component.", "danger")
        return redirect(url_for('main.fee_components'))

    if request.method == 'POST':
        component.name = request.form['name']
        component.description = request.form['description']
        # School_id is not changeable after component creation.
        db.session.commit()
        flash('Fee component updated successfully!', 'success')
        return redirect(url_for('main.fee_components'))

    schools = School.query.filter_by(id=current_user.school_id).all()
    return render_template('edit_fee_component.html', component=component, schools=schools)


@main.route('/delete_fee_component/<int:id>', methods=['POST'])
@login_required
def delete_fee_component(id):
    component = FeeComponent.query.get_or_404(id)
    if component.school_id != current_user.school_id:
        flash("You don't have permission to delete this fee component.", "danger")
        return redirect(url_for('main.fee_components'))
    db.session.delete(component)
    db.session.commit()
    flash('Fee component deleted successfully!', 'success')
    return redirect(url_for('main.fee_components'))

# @main.route('/student_fees')
# def student_fees():
#     # Fetch all student fees
#     fees = StudentFee.query.all()
#     return render_template('view_fees.html', fees=fees)

    # fees = StudentFee.query.filter_by(school_id=current_user.school_id).all()

# @main.route('/student_fees')
# @login_required
# def student_fees():
#     # Ensure the user is linked to a school
#     if not current_user.school_id:
#         flash("No school linked to your account.", "danger")
#         return redirect(url_for('main.index'))

#     school_id = current_user.school_id  # Get logged-in user's school

#     # Fetch only student fees related to the user's school
#     fees = StudentFee.query.join(Student).filter(Student.school_id == school_id).all()

#     return render_template('view_fees.html', fees=fees)

@main.route('/view_class_fees', methods=['GET'])
@login_required
def view_class_fees():
    class_id = request.args.get('class_id')
    academic_year = request.args.get('academic_year')
    term = request.args.get('term')

    classes = Class.query.filter_by(school_id=current_user.school_id).all()
    selected_class = None
    fees = None

    if class_id and academic_year and term:
        class_id = int(class_id)
        selected_class = Class.query.get(class_id)

        if selected_class and selected_class.school_id == current_user.school_id:
            fees = ClassFeeComponent.query.join(FeeComponent).filter(
                ClassFeeComponent.class_id == class_id,
                FeeComponent.school_id == current_user.school_id,
                FeeComponent.academic_year == academic_year, # Filter by academic year
                FeeComponent.term == term # Filter by term
            ).all()

        else:
            flash("Invalid class selection.", "danger")
            return redirect(url_for('main.view_fees'))

    academic_years = [
        year.academic_year
        for year in db.session.query(FeeComponent.academic_year).filter_by(school_id=current_user.school_id).distinct().all()
    ]
    return render_template('view_fees.html',
                           classes=classes,
                           fees=fees,
                           selected_class=selected_class,
                           selected_class_id=class_id,
                           selected_academic_year=academic_year,
                           selected_term=term,
                           academic_years=academic_years)


# @main.route('/generate_class_fees_pdf', methods=['GET'])
# @login_required
# def generate_class_fees_pdf():
#     class_id = request.args.get('class_id')
#     academic_year = request.args.get('academic_year')
#     term = request.args.get('term')

#     if not class_id or not academic_year or not term:
#         flash("Please select Class, Academic Year, and Term.", "danger")
#         return redirect(url_for('main.generate_class_fees'))
    
#     school = School.query.filter_by(id=current_user.school_id).first() # Use current_user's school

#     if not school:
#         return "School information not found", 404

#     # assessment = Assessment.query.filter_by(school_id=school.id).first() # Filter by school

#     # if not assessment:
#     #     return "Assessment information not found", 404

#     school_name = school.name
#     school_address = school.address
#     academic_year = academic_year

#     classes = Class.query.filter_by(school_id=school.id, id=class_id).all() # Filter classes by school and class_id

#     if not classes:
#         flash("Class not found or does not belong to your school.", "danger")
#         return redirect(url_for('main.generate_class_fees'))

#     # classes = Class.query.filter_by(school_id=school.id).all() # Filter classes by school

#     class_fees = []
#     for cls in classes:
#         fee_components = ClassFeeComponent.query.join(FeeComponent).filter(
#             ClassFeeComponent.class_id == cls.id,
#             FeeComponent.school_id == current_user.school_id,
#             FeeComponent.academic_year == academic_year,
#             FeeComponent.term == term
#         ).all()
#         total_fee = sum([component.amount for component in fee_components])
#         class_fees.append({
#             "class_name": cls.class_name,
#             "total_fee": total_fee,
#             "fee_breakdown": [(comp.fee_component.name, comp.amount) for comp in fee_components]
#         })

#     # Generate the PDF
#         # pdf.build(elements)
#         # buffer.seek(0)
#         # return send_file(buffer, as_attachment=True, download_name="class_fees.pdf", mimetype='application/pdf')

#     buffer = BytesIO()
#     pdf = SimpleDocTemplate(buffer, pagesize=letter)
#     elements = []

#     # Add School Header
#     elements.append(Table(
#         [[school_name], [school_address], [academic_year]],
#         style=TableStyle([
#             ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
#             ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
#             ('FONTSIZE', (0, 0), (-1, -1), 14),
#             ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
#         ]),
#         colWidths=[500]
#     ))

#     elements.append(Table([[" "]]))  # Spacer

#     # Add Class Fees Table
#     data = [["Class Name", "Total Fee (₦)", "Fee Breakdown"]]
#     for class_fee in class_fees:
#         fee_breakdown_text = "\n".join([f"{name}: ₦{amount}" for name, amount in class_fee['fee_breakdown']])
#         data.append([class_fee['class_name'], f"₦{class_fee['total_fee']}", fee_breakdown_text])

#     table = Table(data, colWidths=[150, 100, 250])
#     table.setStyle(TableStyle([
#         ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
#         ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
#         ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
#         ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
#         ('FONTSIZE', (0, 0), (-1, 0), 12),
#         ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
#         ('GRID', (0, 0), (-1, -1), 1, colors.black),
#     ]))
#     elements.append(table)

#     pdf.build(elements)

#     # Return the PDF as a downloadable file
#     # buffer.seek(0)
#     # return send_file(buffer, as_attachment=True, download_name="class_fees.pdf", mimetype='application/pdf')

#     pdf.build(elements)
#     buffer.seek(0)
#     return send_file(buffer, as_attachment=True, download_name="class_fees.pdf", mimetype='application/pdf')


@main.route('/generate_class_fees_pdf', methods=['GET'])
@login_required
def generate_class_fees_pdf():
    class_id = request.args.get('class_id')
    academic_year = request.args.get('academic_year')
    term = request.args.get('term')

    if not class_id or not academic_year or not term:
        flash("Please select Class, Academic Year, and Term.", "danger")
        return redirect(url_for('main.generate_class_fees'))

    try:  # Add a try-except block for debugging
        school = School.query.filter_by(id=current_user.school_id).first()

        if not school:
            return "School information not found", 404

        school_name = school.name
        school_address = school.address

        classes = Class.query.filter_by(school_id=school.id, id=class_id).all()

        if not classes:
            flash("Class not found or does not belong to your school.", "danger")
            return redirect(url_for('main.generate_class_fees'))

        class_fees = []
        for cls in classes:
            fee_components = ClassFeeComponent.query.join(FeeComponent).filter(
                ClassFeeComponent.class_id == cls.id,
                FeeComponent.school_id == current_user.school_id,
                FeeComponent.academic_year == academic_year,  # Use the passed academic_year
                FeeComponent.term == term  # Use the passed term
            ).all()

            total_fee = sum(comp.amount for comp in fee_components)
            class_fees.append({
                "class_name": cls.class_name,
                "total_fee": total_fee,
                "fee_breakdown": [(comp.fee_component.name, comp.amount) for comp in fee_components]
            })

        buffer = BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []

        elements.append(Table(
            [[school_name], [school_address], [academic_year]],  # Use academic_year here
            style=TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 14),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]),
            colWidths=[500]
        ))

        elements.append(Table([[" "]]))  # Spacer

        data = [["Class Name", "Total Fee (₦)", "Fee Breakdown"]]
        for class_fee in class_fees:
            fee_breakdown_text = "\n".join([f"{name}: ₦{amount}" for name, amount in class_fee['fee_breakdown']])
            data.append([class_fee['class_name'], f"₦{class_fee['total_fee']}", fee_breakdown_text])

        table = Table(data, colWidths=[150, 100, 250])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)

        pdf.build(elements)  # Only build the PDF once

        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="class_fees.pdf", mimetype='application/pdf')

    except Exception as e:  # Catch any exceptions for debugging
        print(f"Error generating PDF: {e}")  # Print the error message
        flash(f"An error occurred while generating the PDF: {e}", "danger")  # Flash the error
        return redirect(url_for('main.generate_class_fees'))  # Redirect back to the form
    
# @main.route('/generate_fees', methods=['GET']) # Changed to GET to render a template for selecting class, academic year and term
# @login_required
# def generate_fees():
#     schools = School.query.filter_by(id=current_user.school_id).all()
#     classes = Class.query.filter_by(school_id=current_user.school_id).all()

#     # Fetch academic years from the database (replace AcademicYear with your model name)
#     academic_years = [year.academic_year for year in FeeComponent.query.distinct(FeeComponent.academic_year).all()]

#     return render_template('generate_fees.html', schools=schools, classes=classes, academic_years=academic_years) # Pass schools and classes to the template


@main.route('/generate_fees', methods=['GET'])
@login_required
def generate_fees():
    school = School.query.filter_by(id=current_user.school_id).first()

    if not school:
        flash("No school assigned to your account.", "danger")
        return redirect(url_for('main.index'))

    classes = Class.query.filter_by(school_id=school.id).all()

    # Fetch distinct academic years from ClassFeeComponent, filtered by school
    academic_years = [
        year.academic_year
        for year in db.session.query(FeeComponent.academic_year).filter_by(school_id=current_user.school_id).distinct().all()
    ]

    return render_template(
        'generate_fees.html',
        school=school,
        classes=classes,
        academic_years=academic_years
    )
    

@main.route('/generate_class_fees', methods=['GET'])
@login_required
def generate_class_fees():
    classes = Class.query.filter_by(school_id=current_user.school_id).all() # Fetch classes
    return render_template('generate_class_fees.html', selected_academic_year=request.args.get('academic_year'), selected_term=request.args.get('term'), selected_class_id=request.args.get('class_id'), classes = classes) # Pass selected values to the template, include classes


# @main.route('/generate_fees', methods=['GET', 'POST'])
# def generate_fees():
#     if request.method == 'POST':
#         student_id = request.form['student_id']
#         component_id = request.form['component_id']
#         amount = float(request.form['amount'])
#         academic_year = request.form['academic_year']
#         term = request.form['term']

#         fee = StudentFee(
#             student_id=student_id,
#             component_id=component_id,
#             amount=amount,
#             academic_year=academic_year,
#             term=term,
#             payment_status='unpaid'
#         )
#         db.session.add(fee)
#         db.session.commit()
#         return redirect(url_for('main.view_fees', student_id=student_id))
#     return render_template('generate_fees.html')

# @main.route('/view_fees', methods=['GET'])
# def view_fees():
#     student_id = request.args.get('student_id')
#     fees = []
#     if student_id:
#         fees = StudentFee.query.filter_by(student_id=student_id).all()
#     return render_template('view_fees.html', fees=fees, student_id=student_id)


@main.route('/class_fee_components', methods=['GET'])
def display_class_fee_components():
    # Fetch all records from ClassFeeComponent
    class_fee_components = ClassFeeComponent.query.all()

    # Group fees by class
    grouped_fees = defaultdict(list)
    for component in class_fee_components:
        grouped_fees[component.class__.class_name].append(component)

    # Render the template and pass the grouped data
    return render_template(
        'class_fee_components.html', 
        grouped_fees=grouped_fees
    )

# @main.route('/view_class_fees/<int:class_id>')
# def view_class_fees(class_id):
#     # Fetch the specific class by its ID
#     selected_class = Class.query.get_or_404(class_id)
    
   
#     class_fee_components = ClassFeeComponent.query.filter_by(class_id=class_id).all()
#     return render_template('view_class_fees.html', selected_class=selected_class, class_fee_components=class_fee_components)



@main.route('/remove_fee_component_from_class/<int:class_fee_component_id>', methods=['POST'])
def remove_fee_component_from_class(class_fee_component_id):
    class_fee_component = ClassFeeComponent.query.get_or_404(class_fee_component_id)
    class_id = class_fee_component.class_id
    db.session.delete(class_fee_component)
    db.session.commit()
    flash('Fee component removed from class successfully!', 'success')
    return redirect(url_for('main.view_class_fees', class_id=class_id))

@main.route('/record_payment', methods=['GET', 'POST'])
def record_payment():
    if request.method == 'POST':
        student_fee_id = request.form['student_fee_id']
        amount_paid = float(request.form['amount_paid'])
        payment_date = request.form['payment_date']
        payment_method = request.form['payment_method']
        receipt_number = request.form.get('receipt_number')
        notes = request.form.get('notes')

        payment = FeePayment(
            student_fee_id=student_fee_id,
            amount_paid=amount_paid,
            payment_date=payment_date,
            payment_method=payment_method,
            receipt_number=receipt_number,
            notes=notes
        )
        db.session.add(payment)

        # Update payment status
        student_fee = StudentFee.query.get(student_fee_id)
        total_paid = db.session.query(db.func.sum(FeePayment.amount_paid)).filter(
            FeePayment.student_fee_id == student_fee_id
        ).scalar() or 0
        if total_paid >= student_fee.amount:
            student_fee.payment_status = 'paid'
        else:
            student_fee.payment_status = 'unpaid'
        db.session.commit()
        return redirect(url_for('main.view_fees', student_id=student_fee.student_id))
    return render_template('record_payment.html')